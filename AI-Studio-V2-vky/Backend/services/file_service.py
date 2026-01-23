import boto3
import uuid
import os
from fastapi import UploadFile, HTTPException
import PyPDF2
import fitz  # PyMuPDF for enhanced PDF processing
from docx import Document
import io
from utils.snowflake_setup import get_snowflake_config
import snowflake.connector
from utils.content_processor import ContentProcessor
from services.session_manager import DocumentSessionManager
import logging

# Configure logging
logger = logging.getLogger(__name__)


class FileService:
    def __init__(self):
        # S3 Setup
        self.s3_client = boto3.client(
            's3',
            aws_access_key_id=os.getenv('AWS_ACCESS_KEY_ID'),
            aws_secret_access_key=os.getenv('AWS_SECRET_ACCESS_KEY'),
            region_name=os.getenv('AWS_REGION')
        )
        self.bucket_name = os.getenv('S3_BUCKET_NAME')

        # Snowflake Setup
        snowflake_config = get_snowflake_config()
        self.snowflake_conn = snowflake.connector.connect(**snowflake_config)

        # Content Processor for enhanced text extraction
        self.content_processor = ContentProcessor()

        # Session Manager for document tracking
        self.session_manager = DocumentSessionManager()

    async def upload_file(self, file: UploadFile, doc_id: str = None):
        """Upload and process a file with comprehensive error handling"""
        try:
            # Validate file
            if not file.filename:
                raise HTTPException(
                    status_code=400, detail="No filename provided")

            # Generate document ID
            if doc_id is None:
                doc_id = str(uuid.uuid4())

            print(f"Processing file: {file.filename}")

            # Read file content with size validation
            MAX_FILE_SIZE = 50 * 1024 * 1024  # 50MB limit
            file_content = await file.read()

            if len(file_content) > MAX_FILE_SIZE:
                raise HTTPException(
                    status_code=400,
                    detail=f"File too large. Maximum size: {MAX_FILE_SIZE // (1024*1024)}MB"
                )

            if len(file_content) == 0:
                raise HTTPException(status_code=400, detail="File is empty")

            # Upload to S3 with error handling
            s3_key = f"documents/{doc_id}/{file.filename}"
            try:
                self.s3_client.put_object(
                    Bucket=self.bucket_name,
                    Key=s3_key,
                    Body=file_content,
                    ContentType=file.content_type,
                    Metadata={
                        'doc_id': doc_id,
                        'original_filename': file.filename
                    }
                )
                print(f"✅ Uploaded to S3: s3://{self.bucket_name}/{s3_key}")
                logger.info(f"File uploaded to S3: {s3_key}")

            except Exception as e:
                logger.error(f"S3 upload failed for {file.filename}: {str(e)}")
                raise HTTPException(
                    status_code=500,
                    detail=f"S3 upload failed: {str(e)}"
                )

            # Extract text with error handling
            try:
                text_content = await self._extract_text(file_content, file.filename, doc_id=doc_id)
                if not text_content or len(text_content.strip()) == 0:
                    raise ValueError("No text content extracted from file")
                print(f"✅ Extracted {len(text_content)} characters of text")
                logger.info(f"Text extraction successful for {file.filename}")

            except Exception as e:
                logger.error(
                    f"Text extraction failed for {file.filename}: {str(e)}")
                # Clean up S3 upload if text extraction fails
                try:
                    self.s3_client.delete_object(
                        Bucket=self.bucket_name, Key=s3_key)
                    print(f"🧹 Cleaned up S3 object after text extraction failure")
                except Exception as cleanup_error:
                    logger.warning(
                        f"Failed to cleanup S3 object: {cleanup_error}")

                raise HTTPException(
                    status_code=500,
                    detail=f"Text extraction failed: {str(e)}"
                )

            # Store in Snowflake with error handling
            try:
                cursor = self.snowflake_conn.cursor()
                cursor.execute("""
                    INSERT INTO documents (doc_id, filename, s3_key, file_size, file_type, text_content, status)
                    VALUES (%s, %s, %s, %s, %s, %s, 'uploaded')
                """, (doc_id, file.filename, s3_key, len(file_content), file.content_type, text_content))

                self.snowflake_conn.commit()
                print(f"✅ Stored metadata in Snowflake")
                logger.info(f"Database record created for {file.filename}")

            except Exception as e:
                logger.error(
                    f"Database operation failed for {file.filename}: {str(e)}")
                # Clean up S3 upload if database operation fails
                try:
                    self.s3_client.delete_object(
                        Bucket=self.bucket_name, Key=s3_key)
                    print(f"🧹 Cleaned up S3 object after database failure")
                except Exception as cleanup_error:
                    logger.warning(
                        f"Failed to cleanup S3 object: {cleanup_error}")

                raise HTTPException(
                    status_code=500,
                    detail=f"Database operation failed: {str(e)}"
                )

            # Register document with session manager
            try:
                content_summary = text_content[:500] + \
                    "..." if len(text_content) > 500 else text_content
                print(
                    f"📝 Adding document to session: {doc_id} - {file.filename}")

                # Create a new session if none exists
                session_id = self.session_manager.create_session()
                self.session_manager.add_document(
                    session_id=session_id,
                    doc_id=doc_id,
                    filename=file.filename,
                    content_summary=content_summary,
                    file_type=file.content_type
                )
                print(f"✅ Registered document with session manager")
                print(
                    f"📊 Session {session_id} now has {self.session_manager.sessions[session_id]['document_count']} documents")
                logger.info(
                    f"Document registered with session manager: {doc_id}")

            except Exception as e:
                logger.error(
                    f"Session management failed for {file.filename}: {str(e)}")
                # Note: Don't clean up S3/DB here as they're already stored
                # Just log the error and continue
                print(f"⚠️ Session management failed: {e}")

            return {
                "success": True,
                "doc_id": doc_id,
                "session_id": session_id,
                "filename": file.filename,
                "s3_url": f"s3://{self.bucket_name}/{s3_key}",
                "text_content": text_content
            }

        except HTTPException:
            # Re-raise HTTP exceptions as-is
            raise
        except Exception as e:
            logger.error(
                f"Unexpected error in upload_file for {file.filename}: {str(e)}")
            raise HTTPException(
                status_code=500,
                detail=f"Unexpected error during file processing: {str(e)}"
            )

    def get_snowflake_connection(self):
        """Get a fresh Snowflake connection with error handling"""
        try:
            snowflake_config = get_snowflake_config()
            if not all(snowflake_config.values()):
                raise ValueError("Incomplete Snowflake configuration")

            connection = snowflake.connector.connect(**snowflake_config)

            # Test the connection
            cursor = connection.cursor()
            cursor.execute("SELECT 1")
            cursor.fetchone()
            cursor.close()

            logger.info("Snowflake connection established successfully")
            return connection

        except Exception as e:
            logger.error(f"Failed to establish Snowflake connection: {str(e)}")
            raise ValueError(f"Snowflake connection failed: {str(e)}")

    async def _extract_text(self, file_content: bytes, filename: str, doc_id: str | None = None) -> str:
        """Extract text from various file types with error handling"""
        try:
            if filename.lower().endswith('.pdf'):
                return self._extract_pdf_text(file_content, doc_id=doc_id)
            elif filename.lower().endswith(('.docx', '.doc')):
                return self._extract_docx_text(file_content)
            elif filename.lower().endswith('.txt'):
                return file_content.decode('utf-8')
            elif filename.lower().endswith(('.png', '.jpg', '.jpeg', '.tiff')):
                return self.content_processor.process_image_content(file_content, filename)
            else:
                raise ValueError(f"Unsupported file type: {filename}")
        except Exception as e:
            logger.error(f"Text extraction failed for {filename}: {str(e)}")
            raise ValueError(
                f"Failed to extract text from {filename}: {str(e)}")

    def _extract_pdf_text(self, file_content: bytes, doc_id: str | None = None) -> str:
        """Extract text from PDF with enhanced table extraction, page numbers, and error handling"""
        try:
            # Use PyMuPDF (fitz) for better page handling and table extraction
            pdf_file = io.BytesIO(file_content)
            doc = fitz.open(stream=pdf_file, filetype="pdf")

            if len(doc) == 0:
                raise ValueError("PDF has no pages")

            extracted_content = []
            total_pages = len(doc)

            print(f"📄 Processing PDF with {total_pages} pages")
            logger.info(f"Processing PDF with {total_pages} pages")

            for page_num, page in enumerate(doc):
                try:
                    # Extract text from page
                    page_text = page.get_text("text")

                    if page_text and page_text.strip():
                        # Add page number header
                        page_header = f"\n--- PAGE {page_num + 1} ---\n"
                        page_content = page_header + page_text.strip()
                        extracted_content.append((page_num + 1, page_content))

                        print(
                            f"✅ Page {page_num + 1}: Extracted {len(page_text)} characters")
                        logger.info(
                            f"Page {page_num + 1}: Extracted {len(page_text)} characters")
                    else:
                        logger.warning(
                            f"Page {page_num + 1} has no extractable text")

                except Exception as e:
                    logger.warning(
                        f"Failed to extract text from page {page_num + 1}: {e}")
                    print(
                        f"⚠️ Page {page_num + 1}: Text extraction failed - {e}")
                    continue

            doc.close()

            if not extracted_content:
                raise ValueError("No text content could be extracted from PDF")

            # Combine all page content with page numbers
            combined_text = ""
            for page_num, page_content in extracted_content:
                combined_text += page_content + "\n\n"

            # Enhanced PDF processing with comprehensive extraction
            try:
                # Use comprehensive PDF extraction with page numbers
                pdf_details = self.content_processor.extract_pdf_with_page_numbers(
                    file_content)

                if pdf_details:
                    # Log detailed extraction results
                    total_pages = len(pdf_details)
                    total_tables = sum(len(page.get('tables', []))
                                       for page in pdf_details)
                    total_figures = sum(len(page.get('figures', []))
                                        for page in pdf_details)

                    print(f"📊 Comprehensive PDF analysis complete:")
                    print(f"   📄 Total pages: {total_pages}")
                    print(f"   📊 Total tables: {total_tables}")
                    print(f"   🖼️ Total figures: {total_figures}")

                    logger.info(
                        f"Comprehensive PDF analysis: {total_pages} pages, {total_tables} tables, {total_figures} figures")

                    # Create comprehensive metadata summary and store if doc_id provided
                    pdf_metadata = self.content_processor.create_pdf_metadata_summary(
                        pdf_details)
                    if doc_id is not None:
                        # Store metadata in session for later retrieval
                        if hasattr(self, 'session_manager') and pdf_metadata:
                            try:
                                # Store PDF metadata in session
                                self.session_manager.store_pdf_metadata(
                                    doc_id, pdf_metadata)
                                print(f"💾 PDF metadata stored in session")
                                logger.info(
                                    f"PDF metadata stored in session for document {doc_id}")
                            except Exception as e:
                                logger.warning(
                                    f"Failed to store PDF metadata in session: {e}")

                    # Enhance text with detailed page information
                    enhanced_text = ""
                    for page_info in pdf_details:
                        page_num = page_info['page_number']
                        page_content = page_info['content']

                        # Add page header with summary
                        page_header = f"\n{'='*80}\n"
                        page_header += f"📄 PAGE {page_num} | 📊 {len(page_info.get('tables', []))} tables | 🖼️ {len(page_info.get('figures', []))} figures\n"
                        page_header += f"{'='*80}\n"

                        enhanced_text += page_header + page_content + "\n\n"

                        # Add table information if present
                        for table in page_info.get('tables', []):
                            table_markdown = self.content_processor._convert_table_to_markdown(
                                table['data'], page_num, table['table_index']
                            )
                            enhanced_text += table_markdown + "\n"

                    combined_text = enhanced_text

                else:
                    # Fallback to basic table extraction
                    tables = self.content_processor.extract_tables_from_pdf(
                        file_content)
                    if tables:
                        combined_text = self.content_processor.enhance_text_with_tables(
                            combined_text, tables)
                        print(
                            f"✅ Extracted {len(tables)} tables from PDF (fallback method)")
                        logger.info(
                            f"Fallback table extraction: {len(tables)} tables found")

            except Exception as e:
                logger.warning(f"Enhanced PDF processing failed: {e}")
                print(f"⚠️ Enhanced PDF processing failed: {e}")
                # Continue with basic text extraction

            # Log final extraction summary
            print(
                f"📊 PDF extraction complete: {len(extracted_content)} pages processed")
            logger.info(
                f"PDF extraction complete: {len(extracted_content)} pages processed")

            return combined_text

        except Exception as e:
            logger.error(f"PDF text extraction failed: {str(e)}")
            raise ValueError(f"Failed to extract text from PDF: {str(e)}")

    def _extract_docx_text(self, file_content: bytes) -> str:
        """Extract text from DOCX with error handling"""
        try:
            doc = Document(io.BytesIO(file_content))

            if not doc.paragraphs:
                raise ValueError("DOCX has no paragraphs")

            text = ""
            for paragraph in doc.paragraphs:
                if paragraph.text:
                    text += paragraph.text + "\n"

            if not text.strip():
                raise ValueError(
                    "No text content could be extracted from DOCX")

            return text

        except Exception as e:
            logger.error(f"DOCX text extraction failed: {str(e)}")
            raise ValueError(f"Failed to extract text from DOCX: {str(e)}")
